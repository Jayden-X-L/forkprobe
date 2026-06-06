from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/shijunluo/agent_future/forkprobe/bridge-internal-share-final-internal-comms-20260604.md")
OUTPUT = Path("/Users/shijunluo/Downloads/Claude-Codex-Bridge-内部分享稿.docx")

BODY_FONT = "PingFang SC"
CODE_FONT = "Menlo"
BLUE = RGBColor(31, 78, 121)
LIGHT_BLUE = RGBColor(46, 116, 181)
MUTED = RGBColor(89, 89, 89)
INK = RGBColor(32, 32, 32)
CODE_FILL = "F4F6F8"


def set_run_font(run, name=BODY_FONT, size=14, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def add_bottom_border(paragraph, color="D9E2F3", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "6")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def style_document(doc: Document):
    section = doc.sections[0]
    # A4 is a better fit for Chinese internal docs; body text remains 14 pt.
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.78)
    section.right_margin = Inches(0.78)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(14)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 17, BLUE, 12, 6),
        ("Heading 2", 16, LIGHT_BLUE, 10, 5),
        ("Heading 3", 15, BLUE, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run("内部分享稿")
    set_run_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(0)
    run = footer.add_run("Claude-Codex Bridge")
    set_run_font(run, size=9, color=MUTED)


def add_title(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.08
    run = p.add_run(title)
    set_run_font(run, size=22, color=BLUE, bold=True)
    add_bottom_border(p)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(10)
    r = sub.add_run("面向产品、运营、设计、技术等职能角色的内部说明稿")
    set_run_font(r, size=12, color=MUTED)


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=14)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_run_font(run, size=14)


def add_code_block(doc: Document, lines: list[str], language: str | None = None):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.14)
        p.paragraph_format.right_indent = Inches(0.08)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.05
        shade_paragraph(p, CODE_FILL)
        run = p.add_run(("  " + line) if line else "  ")
        set_run_font(run, name=CODE_FONT, size=10.5, color=RGBColor(34, 34, 34))
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(3)


def add_formatted_runs(paragraph, text: str):
    # Minimal parser for inline `code` and **bold** used in the source Markdown.
    idx = 0
    code = False
    bold = False
    buf: list[str] = []

    def flush():
        if not buf:
            return
        value = "".join(buf)
        buf.clear()
        run = paragraph.add_run(value)
        if code:
            set_run_font(run, name=CODE_FONT, size=12, color=RGBColor(34, 34, 34), bold=bold)
        else:
            set_run_font(run, size=14, bold=bold)

    while idx < len(text):
        if text.startswith("**", idx):
            flush()
            bold = not bold
            idx += 2
            continue
        if text[idx] == "`":
            flush()
            code = not code
            idx += 1
            continue
        buf.append(text[idx])
        idx += 1
    flush()


def add_mixed_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.10
    add_formatted_runs(p, text)


def add_mixed_bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    add_formatted_runs(p, text)


def build_docx():
    text = SOURCE.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)

    lines = text.splitlines()
    in_code = False
    code_lines: list[str] = []
    code_lang: str | None = None
    title_done = False

    for raw in lines:
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines, code_lang)
                code_lines = []
                code_lang = None
                in_code = False
            else:
                in_code = True
                code_lang = line.strip("`").strip() or None
            continue

        if in_code:
            code_lines.append(line)
            continue

        if not line:
            continue

        if line.startswith("# "):
            add_title(doc, line[2:].strip())
            title_done = True
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("- "):
            add_mixed_bullet(doc, line[2:].strip())
        else:
            add_mixed_paragraph(doc, line)

    if in_code and code_lines:
        add_code_block(doc, code_lines, code_lang)

    doc.core_properties.title = "Claude-Codex Bridge：让第二个模型复查变成固定流程"
    doc.core_properties.subject = "内部分享稿"
    doc.core_properties.author = "Codex"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build_docx()
    print(out)
